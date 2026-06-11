import docx

d = docx.Document(r"C:\Users\HP\Documents\GitHub\projects\SE-Hakari-Bankai\docs\mine\New Microsoft Word Document.docx")

targets = [
    "1. Introduction",
    "1.5 System Dependencies",
    "2. Process Model",
    "2.1 Selected Model",
    "1.4 Main Features",
    "Manage teachers, courses, rooms, sections, and time slots.",
]
for p in d.paragraphs[:80]:
    t = p.text.strip()
    if t in targets:
        print(f"--- {t!r}")
        print(f"  style={p.style.name}, alignment={p.alignment}")
        pf = p.paragraph_format
        print(f"  space_before={pf.space_before}, space_after={pf.space_after}, left_indent={pf.left_indent}")
        numpr = p._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
        print(f"  numPr={'yes' if numpr is not None else 'no'}")
        for r in p.runs[:3]:
            f = r.font
            print(f"  run: bold={f.bold}, size={f.size}, name={f.name}, color={f.color.rgb if f.color and f.color.type else None}")

print()
print("TABLE STYLES:")
seen = set()
for tbl in d.tables:
    s = tbl.style.name if tbl.style else None
    if s not in seen:
        seen.add(s)
        print(" ", s)

print()
print("DOC DEFAULT FONT:")
st = d.styles["Normal"]
print("  ", st.font.name, st.font.size)
