"""Build the perfected UTOS SRS as a copy of the original compound document.

Keeps the original 7-level structure and all embedded figures; inserts the
missing FR/NFR catalog, definitions, references, roadmap, document control,
and a traceability appendix; fixes the exam-timetabling scope contradiction.
"""
import copy
import shutil

import docx
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

SRC = r"C:\Users\HP\Documents\GitHub\projects\SE-Hakari-Bankai\docs\mine\New Microsoft Word Document.docx"
DST = r"C:\Users\HP\Documents\GitHub\projects\SE-Hakari-Bankai\docs\mine\UTOS_SRS_Perfected.docx"

H1_COLOR = RGBColor(0x1F, 0x4E, 0x79)
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)

shutil.copyfile(SRC, DST)
doc = docx.Document(DST)


def find_para(text, startswith=False):
    for p in doc.paragraphs:
        t = p.text.strip()
        if (startswith and t.startswith(text)) or t == text:
            return p
    raise ValueError(f"paragraph not found: {text!r}")


def rewrite_para(p, new_text, size=11, bold=False, color=None):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    run = p.add_run(new_text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def mk_para(text, size=11, bold=False, color=None, before=0, after=4, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    pf = p.paragraph_format
    if before:
        pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    return p


def mk_h1(t):
    return mk_para(t, 16, True, H1_COLOR, before=20, after=10)


def mk_h2(t):
    return mk_para(t, 13, True, H2_COLOR, before=14, after=7)


def mk_h3(t):
    return mk_para(t, 11.5, True, H2_COLOR, before=10, after=5)


def mk_table(rows, widths=None):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        tbl.style = doc.styles["Table Grid1"]
    except KeyError:
        tbl.style = doc.styles["Table Grid"]
    tbl.autofit = False
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            if i == 0:
                run.font.bold = True
            if widths:
                cell.width = widths[j]
    return tbl


# ---------------------------------------------------------------------------
# 1. Bump version line and add document control before "1. Introduction"
# ---------------------------------------------------------------------------
rewrite_para(find_para("Version 1.0"), "Version 1.1  —  Consolidated Software Requirements Specification")

intro_anchor = find_para("1. Introduction")
created = []
created.append(mk_para("Document Control", 13, True, H2_COLOR, before=14, after=7))
created.append(mk_table(
    [
        ["Version", "Date", "Author", "Description"],
        ["1.0", "May 2026", "Muhammad Hammad Shakeel", "Initial compilation of process model, use cases, domain model, DFD, class diagram, SSDs, and operation contracts."],
        ["1.1", "10 June 2026", "Muhammad Hammad Shakeel", "Consolidated SRS: added full FR/NFR catalog, definitions, references, document roadmap, traceability matrix; aligned scope with the exam-timetabling module."],
    ],
    widths=[Cm(2.0), Cm(2.8), Cm(4.4), Cm(7.8)],
))
created.append(mk_para("", after=6))
for el in created:
    node = el._p if isinstance(el, Paragraph) else el._tbl
    intro_anchor._p.addprevious(node)

# ---------------------------------------------------------------------------
# 2. Fix the scope contradiction (exam timetabling is covered by UC16-UC20)
# ---------------------------------------------------------------------------
scope_p = find_para("Explicitly not in scope", startswith=True)
rewrite_para(scope_p, (
    "The system focuses on generating, managing, and adjusting weekly class timetables for the "
    "teachers, rooms, courses, and student sections of an academic department. It also includes an "
    "exam-timetabling module that schedules end-of-term examinations into exam-suitable rooms, "
    "avoids student and room double-booking, and assigns invigilation duties, reusing the same "
    "master data. Supporting capabilities inside the scope are role-based access for six actor "
    "types, change-request handling, timetable versioning and comparison, publishing, in-app "
    "notifications, export and printing, and operational reports with audit logs. Explicitly not "
    "in scope are bus or transport route management, hostel scheduling, individually personalized "
    "timetables for every student, attendance management, fee management, full university ERP "
    "features, and advanced AI prediction."
))

# ---------------------------------------------------------------------------
# 3. Extend 1.4 Main Features with the system-wide and exam features
# ---------------------------------------------------------------------------
last_bullet = find_para("Export or print the final timetable.")
extra_bullets = [
    "Schedule exams without student or room double-booking.",
    "Assign and view invigilation duties for teachers.",
    "Submit, review, approve, or reject change requests.",
    "Keep timetable versions and compare any two versions.",
    "Send in-app notifications for publish, lock, and request events.",
    "Provide role-based dashboards for six actor types.",
    "Record audit logs of who changed what and when.",
]
prev = last_bullet
for text in extra_bullets:
    new_el = copy.deepcopy(last_bullet._p)
    new_p = Paragraph(new_el, last_bullet._parent)
    for r in list(new_p.runs):
        r._r.getparent().remove(r._r)
    run = new_p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    prev._p.addnext(new_el)
    prev = new_p

# ---------------------------------------------------------------------------
# 4. Insert sections 1.6 - 1.10 before "2. Process Model"
# ---------------------------------------------------------------------------
anchor = find_para("2. Process Model")
created = []
add = created.append

# --- 1.6 Definitions ---
add(mk_h2("1.6 Definitions, Acronyms, and Abbreviations"))
add(mk_table(
    [
        ["Term", "Definition"],
        ["UTOS", "University Timetable Optimization System — the system specified by this document."],
        ["SRS", "Software Requirements Specification."],
        ["FR / NFR", "Functional Requirement / Non-Functional Requirement."],
        ["UC", "Use Case (UC00–UC22, specified in Part II)."],
        ["SSD", "System Sequence Diagram (Part VI)."],
        ["OC", "Operation Contract (OC-00–OC-22, Part VII)."],
        ["DFD", "Data Flow Diagram (Part IV)."],
        ["SuD", "System under Discussion — the UTOS web application within its boundary."],
        ["Hard constraint", "A rule the solver must never violate (e.g., no teacher in two rooms at once)."],
        ["Soft preference", "A weighted, desirable rule the solver tries to satisfy and penalizes when violated (e.g., morning classes)."],
        ["Master data", "The validated base records: teachers, rooms, courses, sections, timeslots, holidays."],
        ["Timetable version", "One saved generation or repair result with status Draft, Published, or Archived."],
        ["Change request", "A structured user request to alter a published or draft timetable entry."],
        ["Solver Engine", "The internal optimization component that places class sessions and exams."],
        ["Repair / Re-optimization", "A solver run that fixes a timetable after approved changes while keeping locked entries and minimizing disruption."],
        ["Invigilation", "Supervision duty assigned to a teacher for an examination session."],
        ["RBAC", "Role-Based Access Control — access decided by the user's assigned role."],
        ["MFA", "Multi-Factor Authentication — second login factor, enforced for administrator accounts."],
        ["CRUD", "Create, Read, Update, Delete operations on records."],
        ["MoSCoW", "Priority scheme: Must have, Should have, Could have, Won't have (this release)."],
    ],
    widths=[Cm(4.2), Cm(12.8)],
))
add(mk_para("", after=4))

# --- 1.7 References ---
add(mk_h2("1.7 References"))
refs = [
    "IEEE Std 830-1998 / ISO/IEC/IEEE 29148:2018 — Recommended practice and standard for software requirements specifications (structure and quality criteria followed by this document).",
    "C. Larman, Applying UML and Patterns, 3rd ed. — source of the high-level/expanded use case, domain model, SSD, and operation contract formats used in Parts II, III, VI, and VII.",
    "I. Sommerville, Software Engineering, 10th ed. — incremental process model and requirements engineering guidance used in Part I, Section 2.",
    "R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner's Approach, 9th ed. — process model selection and DFD (Yourdon/DeMarco) notation used in Part IV.",
    "UTOS project repository (SE-Hakari-Bankai) — implementation, database schema, and UI evidence that this specification traces to.",
]
for i, r in enumerate(refs, 1):
    add(mk_para(f"[{i}]  {r}", after=4))

# --- 1.8 Functional Requirements ---
add(mk_h2("1.8 Functional Requirements"))
add(mk_para(
    "This section defines the complete functional requirement catalog for UTOS. Every requirement "
    "is uniquely identified, written as a verifiable \"shall\" statement, and prioritized using the "
    "MoSCoW scheme (M = Must have, S = Should have, C = Could have). The requirement groups FR-00 "
    "through FR-21 are the targets of the cross-references used by the use cases in Part II and "
    "the operation contracts in Part VII; the full mapping is given in Appendix A.", after=8))

FR_GROUPS = [
    ("FR-00 — Authentication and Role-Based Access (UC00)", [
        ("FR-00.1", "The system shall provide a login screen where a user identifies with credentials and an assigned role.", "M"),
        ("FR-00.2", "The system shall validate the supplied credentials against stored user accounts before granting access.", "M"),
        ("FR-00.3", "The system shall create a session on successful login and keep it active until logout or timeout.", "M"),
        ("FR-00.4", "The system shall open the role-specific dashboard immediately after a successful login.", "M"),
        ("FR-00.5", "The system shall restrict every screen, navigation item, and action to the roles permitted to use it.", "M"),
        ("FR-00.6", "The system shall provide a logout function that terminates the session and returns to the login screen.", "M"),
        ("FR-00.6a", "The system shall reject invalid credentials with a plain-language error message and log the attempt.", "M"),
        ("FR-00.6b", "The system shall deny access to a user with no assigned role and direct them to the System Administrator.", "M"),
        ("FR-00.6c", "The system shall automatically log out inactive non-administrator sessions according to the configured security settings.", "S"),
    ]),
    ("FR-01 — Organization Hierarchy (UC01)", [
        ("FR-01.1", "The system shall let the Timetable Administrator create and edit the University → Faculty → Department → Batch → Section hierarchy.", "M"),
        ("FR-01.2", "The system shall let the administrator define academic terms with names and date ranges.", "M"),
        ("FR-01.3", "The system shall validate uniqueness, parent linkage, and naming rules for every hierarchy unit before saving.", "M"),
        ("FR-01.4", "The system shall allow moving or deactivating a unit and shall warn about dependent records before applying the change.", "S"),
        ("FR-01.5", "The system shall block deletion of any unit that has dependent timetable entries until they are reassigned.", "M"),
    ]),
    ("FR-02 — Academic Calendar and Master Data (UC02, UC03)", [
        ("FR-02.1", "The system shall let the administrator define semester start and end dates for each academic term.", "M"),
        ("FR-02.2", "The system shall let the administrator record official holidays, blackout dates, and makeup-class windows.", "M"),
        ("FR-02.3", "The system shall let the administrator configure working days per week, and the solver shall never place a class on a blocked day.", "M"),
        ("FR-02.4", "The system shall validate calendar entries for overlapping ranges and end-before-start errors before saving.", "M"),
        ("FR-02.5", "The system shall provide create, read, update, and delete operations for teachers, rooms, courses, sections, and timeslots.", "M"),
        ("FR-02.6", "The system shall validate master-data records for required fields, duplicates, capacity values, and relationship rules.", "M"),
        ("FR-02.6a", "The system shall block (or require replacement before) deletion of any master-data record used by an existing timetable entry.", "M"),
        ("FR-02.6b", "The system shall show the Department Coordinator and Facility Manager filtered, role-appropriate views of master data.", "S"),
    ]),
    ("FR-03 — Constraints and Preferences (UC04)", [
        ("FR-03.1", "The system shall enforce that no teacher is assigned to two class sessions in the same timeslot (hard).", "M"),
        ("FR-03.2", "The system shall enforce that no room hosts two class sessions in the same timeslot (hard).", "M"),
        ("FR-03.3", "The system shall enforce that no student section attends two class sessions in the same timeslot (hard).", "M"),
        ("FR-03.4", "The system shall enforce that an assigned room's capacity is at least the size of the attending section (hard).", "M"),
        ("FR-03.5", "The system shall enforce required room types (e.g., lab sessions only in laboratory rooms) (hard).", "M"),
        ("FR-03.6", "The system shall exclude holidays and blackout dates from all generated schedules (hard).", "M"),
        ("FR-03.7", "The system shall respect recorded teacher unavailability when placing sessions (hard).", "M"),
        ("FR-03.8", "The system shall enforce each teacher's configured maximum daily teaching load (hard).", "M"),
        ("FR-03.9", "The system shall support a weighted soft preference for scheduling classes in morning timeslots.", "S"),
        ("FR-03.10", "The system shall support a weighted soft preference for ending each day's classes as early as possible.", "S"),
        ("FR-03.11", "The system shall support a weighted soft preference for energy saving by compacting classes into fewer buildings and contiguous slots.", "S"),
        ("FR-03.12", "The system shall support a weighted soft preference for assigning consecutive classes to nearby rooms (proximity).", "S"),
        ("FR-03.13", "The system shall support a weighted soft preference for reducing student movement (traffic) between buildings.", "S"),
        ("FR-03.14", "The system shall let authorized users enable, disable, and weight each soft preference within an allowed range.", "M"),
        ("FR-03.15", "The system shall validate the constraint configuration for completeness and shall warn or block when a hard rule is disabled unsafely.", "M"),
    ]),
    ("FR-04 — Class Timetable Generation (UC05)", [
        ("FR-04.1", "The system shall let the Timetable Administrator trigger timetable generation for a selected academic term.", "M"),
        ("FR-04.2", "The system shall validate, before generation, that the required master data, calendar, and constraint configuration exist.", "M"),
        ("FR-04.3", "The system shall build a solver problem model from teachers, rooms, courses, sections, timeslots, calendar, availability, and active constraints.", "M"),
        ("FR-04.4", "The system shall attempt to assign every class session to a valid room and timeslot satisfying all hard constraints.", "M"),
        ("FR-04.5", "The system shall save each generation result as a new draft timetable version.", "M"),
        ("FR-04.6", "The system shall report the quality score, hard-conflict count, soft penalty, and unplaced session count for every run.", "M"),
        ("FR-04.7", "The system shall show generation progress and complete or abort within a configurable time limit.", "S"),
        ("FR-04.7a", "If master data is missing, the system shall list the missing teachers, rooms, timeslots, courses, or sections and stop generation.", "M"),
        ("FR-04.7b", "If no complete timetable is feasible, the system shall save the partial result and list each unplaced session with the reason.", "M"),
        ("FR-04.7c", "The system shall try alternate room and timeslot assignments before reporting a placement failure.", "M"),
    ]),
    ("FR-05 — Timetable Quality Review (UC06)", [
        ("FR-05.1", "The system shall display each draft timetable as a weekly grid together with a summary status.", "M"),
        ("FR-05.2", "The system shall list every conflict and warning with the affected entries and an explanation of the issue.", "M"),
        ("FR-05.3", "The system shall display the quality score with a breakdown of soft-constraint penalties.", "M"),
        ("FR-05.4", "The system shall display room-utilization and teacher-load statistics for the reviewed version.", "S"),
        ("FR-05.5", "The system shall let reviewers proceed from review to manual adjustment, re-optimization, or publishing.", "M"),
    ]),
    ("FR-06 — Manual Timetable Adjustment (UC07)", [
        ("FR-06.1", "The system shall let authorized users manually move or edit a draft timetable entry.", "M"),
        ("FR-06.2", "The system shall validate every manual change against teacher, room, section, and timeslot hard constraints.", "M"),
        ("FR-06.3", "The system shall reject a move into an occupied slot and display the colliding entry.", "M"),
        ("FR-06.4", "The system shall record every manual change in the audit history (who, what, when).", "M"),
    ]),
    ("FR-07 — Locking and Re-optimization (UC07, UC10)", [
        ("FR-07.1", "The system shall let the administrator lock and unlock individual timetable entries.", "M"),
        ("FR-07.2", "The system shall keep all locked entries unchanged during every re-optimization (repair) run.", "M"),
        ("FR-07.3", "The system shall repair the timetable after approved changes by altering the smallest possible set of entries and shall report a disruption summary.", "M"),
    ]),
    ("FR-08 — Publishing (UC11)", [
        ("FR-08.1", "The system shall let the administrator publish an approved draft as the official timetable of a term.", "M"),
        ("FR-08.2", "The system shall keep exactly one published version per term and archive the previously published version.", "M"),
        ("FR-08.3", "The system shall check conflict status and required approvals before publication and block (or require override authority) when hard conflicts remain.", "M"),
        ("FR-08.4", "The system shall notify all affected teachers and students when a timetable is published.", "M"),
        ("FR-08.4a", "If the administrator cancels publication, the system shall keep the version as a draft with no visible change to users.", "M"),
        ("FR-08.4b", "The system shall make the published timetable visible to each user filtered by their role and identity.", "M"),
    ]),
    ("FR-09 — Timetable Viewing (UC12)", [
        ("FR-09.1", "The system shall show each teacher a personal view of their own teaching schedule from the published version.", "M"),
        ("FR-09.2", "The system shall show each student the schedule of their own section from the published version.", "M"),
        ("FR-09.3", "The system shall display each entry with day, time, course, room, and counterpart (teacher or section).", "M"),
        ("FR-09.4", "The system shall show a clear empty-state message when no published timetable exists or the user has no assignment.", "M"),
    ]),
    ("FR-10 — Change Requests (UC08, UC09)", [
        ("FR-10.1", "The system shall let a teacher or coordinator submit a structured change request with subject, summary, reason, and an optional attachment, recorded with status Pending.", "M"),
        ("FR-10.2", "The system shall let an authorized reviewer inspect a pending request with conflict-check support and approve or reject it with an optional comment.", "M"),
        ("FR-10.3", "The system shall notify the requester of the decision and route approved requests to re-optimization or manual adjustment.", "M"),
    ]),
    ("FR-11 — Timetable Versioning (UC05, UC10, UC15)", [
        ("FR-11.1", "The system shall create a new timetable version for every generation and repair run instead of overwriting earlier results.", "M"),
        ("FR-11.2", "The system shall track each version's status as Draft, Published, or Archived.", "M"),
        ("FR-11.3", "The system shall store version metadata: name, creation time, quality score, conflict counts, and unplaced count.", "M"),
    ]),
    ("FR-12 — Export and Printing (UC13)", [
        ("FR-12.1", "The system shall export a selected timetable to PDF, Excel, and CSV formats.", "S"),
        ("FR-12.2", "The system shall produce the PDF export in a printable A4 layout with day, time, course, teacher, room, and section.", "S"),
        ("FR-12.3", "The system shall let the user choose the export scope (full, personal, or department) within their permissions.", "S"),
        ("FR-12.4", "The system shall report export failures with an error message and a retry option.", "S"),
    ]),
    ("FR-13 — Notifications (UC14)", [
        ("FR-13.1", "The system shall create an in-app notification when generation completes, a version is created, an entry is locked or unlocked, a request decision is made, or a timetable is published.", "S"),
        ("FR-13.2", "The system shall route each notification only to the users affected by the event.", "S"),
        ("FR-13.3", "The system shall show each user a notification list with an unread-count badge.", "S"),
        ("FR-13.4", "The system shall mark a notification as read when opened and navigate to the related screen.", "S"),
        ("FR-13.5", "The system shall let users disable notification categories; suppressed notifications are stored but not surfaced as unread.", "C"),
    ]),
    ("FR-14 — Version Comparison (UC15)", [
        ("FR-14.1", "The system shall let an authorized user select two timetable versions of the same term for comparison.", "S"),
        ("FR-14.2", "The system shall compute and display the differences as added, removed, and changed entries with totals.", "S"),
        ("FR-14.3", "The system shall block comparison of versions belonging to different terms with a clear warning.", "S"),
    ]),
    ("FR-15 — Exam Room Management (UC16)", [
        ("FR-15.1", "The system shall reuse class-room master data as the base for exam room management.", "M"),
        ("FR-15.2", "The system shall let the administrator and facility manager flag rooms as exam-suitable with a separate examination capacity.", "M"),
        ("FR-15.3", "The system shall record exam-window availability per room and flag maintenance conflicts with planned exam days.", "S"),
        ("FR-15.4", "The system shall warn when an entered exam capacity exceeds the room's class capacity and require confirmation.", "S"),
    ]),
    ("FR-16 — Student Hierarchy (UC17)", [
        ("FR-16.1", "The system shall maintain the student hierarchy down to individual students within sections.", "M"),
        ("FR-16.2", "The system shall support bulk import of students and enrollments, reporting and skipping invalid rows by row number.", "M"),
        ("FR-16.3", "The system shall maintain section-course enrollment links used to detect exam clashes.", "M"),
        ("FR-16.4", "The system shall warn when a student is enrolled in conflicting sections.", "S"),
    ]),
    ("FR-17 — Exam Course Management (UC18)", [
        ("FR-17.1", "The system shall let the administrator mark which courses have exams in a given exam session.", "M"),
        ("FR-17.2", "The system shall record an exam duration per exam course and validate it against the allowed slot length.", "M"),
        ("FR-17.3", "The system shall link each exam course to the sections that will sit the exam.", "M"),
        ("FR-17.4", "The system shall recalculate the expected student count per exam course so room capacity can be matched.", "M"),
    ]),
    ("FR-18 — Exam Scheduling (UC19)", [
        ("FR-18.1", "The system shall load exam courses, enrolled sections, students, exam-suitable rooms, and exam rules before scheduling.", "M"),
        ("FR-18.2", "The system shall suggest room and time placements for every exam in the exam window.", "M"),
        ("FR-18.3", "The system shall never schedule one student in two exams at the same time (hard).", "M"),
        ("FR-18.4", "The system shall never double-book an exam room (hard).", "M"),
        ("FR-18.5", "The system shall match exam enrollment to room capacity and suggest splitting across rooms or using larger rooms when capacity is insufficient.", "M"),
        ("FR-18.6", "The system shall limit each student to at most two exams per day.", "M"),
        ("FR-18.7", "The system shall propose invigilation assignments based on teacher availability.", "S"),
        ("FR-18.8", "The system shall let the administrator confirm or modify suggestions and save the exam schedule.", "M"),
    ]),
    ("FR-19 — Exam Schedule Viewing (UC20)", [
        ("FR-19.1", "The system shall show each student a personal list of their exams with date, time, course, and room.", "M"),
        ("FR-19.2", "The system shall show each teacher their invigilation duties.", "M"),
        ("FR-19.3", "The system shall show the facility manager room-level exam occupation.", "S"),
        ("FR-19.4", "The system shall support exporting and printing the exam views.", "S"),
        ("FR-19.5", "The system shall show clear empty-state notices when no exams or duties exist for the user.", "M"),
    ]),
    ("FR-20 — User and Role Management (UC21)", [
        ("FR-20.1", "The system shall let the System Administrator create, edit, and deactivate user accounts for all six roles.", "M"),
        ("FR-20.2", "The system shall let the System Administrator assign and revoke roles, updating the user's effective permissions immediately.", "M"),
        ("FR-20.3", "The system shall support bulk import of user accounts, importing valid rows and reporting invalid ones.", "S"),
        ("FR-20.4", "The system shall support password reset via reset link or temporary password.", "M"),
        ("FR-20.5", "The system shall block removal of the last administrator role and duplicate email addresses.", "M"),
    ]),
    ("FR-21 — Reports and Audit Logs (UC22)", [
        ("FR-21.1", "The system shall provide a room-utilization report with capacity warnings, peak periods, and free rooms.", "S"),
        ("FR-21.2", "The system shall provide a teacher-load report with daily load and balance metrics.", "S"),
        ("FR-21.3", "The system shall provide conflict summaries per timetable version.", "S"),
        ("FR-21.4", "The system shall provide a student-gap analysis counting free periods between classes per section.", "C"),
        ("FR-21.5", "The system shall display quality-score trends across versions of a term.", "C"),
        ("FR-21.6", "The system shall provide an energy-usage indication based on building and slot compaction.", "C"),
        ("FR-21.7", "The system shall record an audit-log entry (action, timestamp, user, old value, new value) for every state-changing operation.", "M"),
        ("FR-21.8", "The system shall let authorized users filter reports by period, building, department, or term, and export the results.", "S"),
    ]),
]

FR_WIDTHS = [Cm(2.2), Cm(12.4), Cm(2.4)]
for title, rows in FR_GROUPS:
    add(mk_h3(title))
    add(mk_table([["ID", "Requirement", "Priority"]] + list(rows), widths=FR_WIDTHS))
    add(mk_para("", after=4))

# --- 1.9 Non-Functional Requirements ---
add(mk_h2("1.9 Non-Functional Requirements"))
add(mk_para(
    "Non-functional requirements constrain how the system delivers the functions above. "
    "Each NFR is verifiable through the stated target. NFR-05 (security) and NFR-07 (audit) "
    "are cross-referenced from the use cases in Part II.", after=8))
add(mk_table(
    [
        ["ID", "Category", "Requirement and Target"],
        ["NFR-01", "Performance", "Timetable generation for one department (up to 50 courses, 20 teachers, 15 rooms, 40 weekly timeslots) shall complete within 60 seconds; all interactive pages shall respond within 2 seconds under normal load."],
        ["NFR-02", "Usability", "A new user shall be able to view their timetable within 2 minutes of first login without training; all error messages shall be in plain language and state the corrective action."],
        ["NFR-03", "Reliability & Availability", "The system shall be available 99% of the time during the academic semester; a failed generation run shall never corrupt or overwrite an existing timetable version."],
        ["NFR-04", "Scalability", "The data model and solver interface shall support growth to at least 5 departments, 200 teachers, 100 rooms, and 5,000 students without architectural change."],
        ["NFR-05", "Security", "Access to every function shall be controlled by role (RBAC); passwords shall be stored as salted hashes; MFA shall be enforced for administrator accounts; sessions shall expire after configurable inactivity."],
        ["NFR-06", "Maintainability & Portability", "The solver shall be replaceable behind a stable solve(problem) → result interface without changes to other layers; the web client shall run on current Chrome, Edge, and Firefox without plugins."],
        ["NFR-07", "Audit & Accountability", "Every state-changing operation shall produce an immutable audit record (actor, action, timestamp, old/new values) retained for at least one academic year."],
        ["NFR-08", "Data Integrity & Backup", "All writes shall be transactional; the database shall be backed up daily; restoring a backup shall lose at most 24 hours of data (RPO ≤ 24 h)."],
    ],
    widths=[Cm(1.8), Cm(3.6), Cm(11.6)],
))
add(mk_para("", after=4))

# --- 1.10 Document Roadmap ---
add(mk_h2("1.10 Document Roadmap — The Seven Levels"))
add(mk_para(
    "This specification is organized as seven analysis levels, in the order they were produced. "
    "Each level refines the previous one: the process model frames how the work is delivered, the "
    "requirements and use cases say what the system must do, the domain model and DFD describe the "
    "problem domain and data movement, and the class diagram, system sequence diagrams, and "
    "operation contracts make the model precise enough to design and test against. Appendix A "
    "traces every use case to its requirements and operation contract.", after=8))
add(mk_table(
    [
        ["Level", "Artifact", "Where", "Purpose"],
        ["1", "Process Model + Requirements", "Part I (Sections 1–2)", "Scope, FR/NFR catalog, and the incremental delivery plan with its Agile/Scrum fallback."],
        ["2", "Use Case Analysis", "Part II (UC00–UC22)", "Actor goals as high-level and fully dressed expanded use cases with alternative courses."],
        ["3", "Domain Model", "Part III (Assignment 03)", "Real-world concepts, attributes, and associations — no software artifacts."],
        ["4", "Data Flow Diagrams", "Part IV (Assignment 04)", "Level 0 context, functional hierarchy, Level 1, and Level 2 decompositions with balancing."],
        ["5", "Class Diagram", "Part V (Assignment 05)", "Design classes, attributes, methods, relationships, multiplicities, and enumerations."],
        ["6", "System Sequence Diagrams", "Part VI", "Actor–system interactions for the main success scenario of every use case."],
        ["7", "Operation Contracts", "Part VII (OC-00–OC-22)", "Pre/postconditions, outputs, and exceptions for every system operation."],
    ],
    widths=[Cm(1.4), Cm(4.6), Cm(4.0), Cm(7.0)],
))
add(mk_para("", after=8))

for el in created:
    node = el._p if isinstance(el, Paragraph) else el._tbl
    anchor._p.addprevious(node)

# ---------------------------------------------------------------------------
# 5. Retitle the part-1 end marker
# ---------------------------------------------------------------------------
end_p = find_para("End of Document", startswith=True)
rewrite_para(end_p, "End of Part I — Introduction, Requirements, and Process Model  ·  UTOS v1.1",
             size=10)

# ---------------------------------------------------------------------------
# 6. Append Appendix A — Requirements Traceability Matrix
# ---------------------------------------------------------------------------
brk = doc.add_paragraph()
brk.add_run().add_break(WD_BREAK.PAGE)

mk_h1("Appendix A — Requirements Traceability Matrix")
mk_para(
    "The matrix below traces every use case (Part II) to the functional requirements it realizes "
    "(Section 1.8), the operation contract that formalizes its system operation (Part VII), and the "
    "delivery increment from the process model (Section 2.4). Exam-timetabling use cases are "
    "delivered as Phase 2, after Increment 6 of the class-timetabling system. Every FR group is "
    "covered by at least one use case and every use case has exactly one operation contract, so "
    "coverage is complete in both directions.", after=8)

trace_rows = [
    ["Use Case", "Title", "Functional Requirements", "Contract", "Increment"],
    ["UC00", "Authenticate and Login", "FR-00.1–FR-00.6c; NFR-05", "OC-00", "1 (used by all)"],
    ["UC01", "Maintain Organization Hierarchy", "FR-01.1–FR-01.5", "OC-01", "1"],
    ["UC02", "Maintain Academic Calendar", "FR-02.1–FR-02.4", "OC-02", "1"],
    ["UC03", "Manage Master Data", "FR-02.5–FR-02.6b; FR-15.1", "OC-03", "1"],
    ["UC04", "Define Constraints and Preferences", "FR-03.1–FR-03.15", "OC-04", "2"],
    ["UC05", "Generate Class Timetable", "FR-04.1–FR-04.7c; FR-11.1–FR-11.3", "OC-05", "3"],
    ["UC06", "Review Timetable Quality", "FR-05.1–FR-05.5; FR-21.1–FR-21.3", "OC-06", "3"],
    ["UC07", "Manually Adjust and Lock Timetable", "FR-06.1–FR-06.4; FR-07.1", "OC-07", "4"],
    ["UC08", "Submit Change Request", "FR-10.1; FR-13.1–FR-13.2", "OC-08", "5"],
    ["UC09", "Approve or Reject Change Request", "FR-10.2–FR-10.3", "OC-09", "5"],
    ["UC10", "Re-optimize Timetable", "FR-07.2–FR-07.3; FR-11.1", "OC-10", "5"],
    ["UC11", "Publish Timetable", "FR-08.1–FR-08.4b", "OC-11", "6"],
    ["UC12", "View Personal or Section Timetable", "FR-09.1–FR-09.4", "OC-12", "6"],
    ["UC13", "Export Class Timetable", "FR-12.1–FR-12.4", "OC-13", "6"],
    ["UC14", "Manage Notifications", "FR-13.1–FR-13.5", "OC-14", "6"],
    ["UC15", "Compare Timetable Versions", "FR-14.1–FR-14.3; FR-11.2–FR-11.3", "OC-15", "5"],
    ["UC16", "Manage Exam Rooms", "FR-15.1–FR-15.4", "OC-16", "Phase 2"],
    ["UC17", "Manage Student Hierarchy", "FR-16.1–FR-16.4", "OC-17", "Phase 2"],
    ["UC18", "Manage Exam Courses", "FR-17.1–FR-17.4", "OC-18", "Phase 2"],
    ["UC19", "Schedule Exams", "FR-18.1–FR-18.8", "OC-19", "Phase 2"],
    ["UC20", "View Exam Schedule and Invigilation Duties", "FR-19.1–FR-19.5", "OC-20", "Phase 2"],
    ["UC21", "Manage Users and Roles", "FR-20.1–FR-20.5; NFR-05", "OC-21", "1"],
    ["UC22", "View System Reports and Audit Logs", "FR-21.1–FR-21.8; NFR-07", "OC-22", "6"],
]
mk_table(trace_rows, widths=[Cm(1.7), Cm(6.0), Cm(5.4), Cm(1.9), Cm(2.0)])
mk_para("", after=6)
mk_para("Each system sequence diagram in Part VI carries the same UC number as its use case "
        "(Figure UC00 … Figure UC22), so the trace extends Use Case → SSD → Operation Contract "
        "without renumbering.", after=4)
mk_para("End of Document  ·  University Timetable Optimization System  ·  SRS v1.1", 10, after=4)

doc.save(DST)
print("saved:", DST)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
