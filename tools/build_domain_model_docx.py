from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets" / "domain_model_assignment"
OUT = DOCS / "Assignment_03_Domain_Model_UTOS.docx"
SOURCE_DIAGRAMS = DOCS / "proper docs for srs building"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GREEN = "2F6F4E"
LIGHT_GREEN = "E4F1E7"
GRAY = "F4F6F8"
CHARCOAL = "222222"
MUTED = "666666"
ACCENT = "B33A3A"


def font(size=32, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap_text(draw, text, fnt, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        bbox = draw.textbbox((0, 0), trial, font=fnt)
        if bbox[2] - bbox[0] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, xy, text, fnt, fill, max_width, line_gap=7):
    x, y = xy
    for line in wrap_text(draw, text, fnt, max_width):
        draw.text((x, y), line, font=fnt, fill=fill)
        bbox = draw.textbbox((0, 0), line, font=fnt)
        y += (bbox[3] - bbox[1]) + line_gap
    return y


def box(draw, xy, title, body=None, fill="#FFFFFF", outline="#1F4E79", title_fill="#1F4E79"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=4)
    draw.text((x1 + 24, y1 + 18), title, font=font(31, True), fill=title_fill)
    if body:
        draw_wrapped(draw, (x1 + 24, y1 + 62), body, font(22), "#222222", x2 - x1 - 48, 6)


def center_text(draw, xy, text, fnt, fill):
    x1, y1, x2, y2 = xy
    bbox = draw.textbbox((0, 0), text, font=fnt)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=fnt, fill=fill)


def line(draw, start, end, label="", color="#444444"):
    draw.line([start, end], fill=color, width=4)
    if label:
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        bbox = draw.textbbox((0, 0), label, font=font(19, True))
        pad = 8
        draw.rounded_rectangle(
            (mx - bbox[2] / 2 - pad, my - 18, mx + bbox[2] / 2 + pad, my + 18),
            radius=8,
            fill="#FFFFFF",
            outline="#DDDDDD",
        )
        draw.text((mx - bbox[2] / 2, my - 12), label, font=font(19, True), fill=color)


def create_method_image(path):
    img = Image.new("RGB", (1800, 740), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 740), fill="#FFFFFF")
    draw.text((70, 48), "Domain Modeling Method Used", font=font(54, True), fill=f"#{BLUE}")
    draw.text(
        (70, 120),
        "Based on lecture rule: identify real-world concepts first, then attributes and associations.",
        font=font(26),
        fill=f"#{MUTED}",
    )

    steps = [
        ("1", "Find Concepts", "Use concept categories and noun phrases from UTOS requirements."),
        ("2", "Add Attributes", "Keep simple data values only: Text, Number, Date, Time, Boolean."),
        ("3", "Add Associations", "Preserve need-to-know relationships, not database foreign keys."),
        ("4", "Add Multiplicity", "Show 1, 0..1, 0..*, and 1..* at association ends."),
    ]
    x = 80
    y = 245
    w = 390
    h = 300
    for idx, title, body in steps:
        box(draw, (x, y, x + w, y + h), title, body, fill="#F8FBFD", outline=f"#{BLUE}")
        draw.ellipse((x + 22, y + h - 82, x + 86, y + h - 18), fill=f"#{GREEN}")
        center_text(draw, (x + 22, y + h - 82, x + 86, y + h - 18), idx, font(28, True), "#FFFFFF")
        if idx != "4":
            draw.line((x + w + 20, y + h / 2, x + w + 82, y + h / 2), fill=f"#{ACCENT}", width=5)
            draw.polygon(
                [
                    (x + w + 82, y + h / 2),
                    (x + w + 58, y + h / 2 - 15),
                    (x + w + 58, y + h / 2 + 15),
                ],
                fill=f"#{ACCENT}",
            )
        x += 430

    draw.rounded_rectangle((70, 610, 1730, 690), radius=20, fill="#FFF4E8", outline="#E0A75E", width=3)
    draw.text(
        (105, 632),
        "Excluded from model: screens, database tables, API routes, methods, solver internals, and foreign-key attributes.",
        font=font(28, True),
        fill="#7A4B00",
    )
    img.save(path)


def create_domain_diagram(path):
    img = Image.new("RGB", (2200, 1600), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "UTOS Domain Model Overview", font=font(58, True), fill=f"#{BLUE}")
    draw.text(
        (70, 115),
        "Concepts, simple attributes, and need-to-know associations for the class timetable domain.",
        font=font(27),
        fill=f"#{MUTED}",
    )

    draw.rounded_rectangle((55, 185, 2145, 760), radius=20, fill="#FBFDFF", outline="#D5E5F0", width=3)
    draw.text((90, 205), "Academic Master Data", font=font(34, True), fill=f"#{GREEN}")
    draw.rounded_rectangle((55, 835, 2145, 1490), radius=20, fill="#FCFDFB", outline="#DDE8D9", width=3)
    draw.text((90, 855), "Scheduling and Change Domain", font=font(34, True), fill=f"#{GREEN}")

    boxes = {
        "University": (95, 300, 375, 415),
        "Faculty": (455, 300, 735, 415),
        "Department": (815, 300, 1120, 415),
        "Teacher": (1250, 245, 1560, 375),
        "StudentSection": (1250, 425, 1560, 575),
        "Student": (1660, 425, 1950, 555),
        "Course": (815, 610, 1120, 730),
        "AcademicTerm": (95, 560, 400, 700),
        "Timeslot": (470, 535, 760, 655),
        "Holiday": (470, 665, 760, 735),
        "Building": (1250, 610, 1530, 730),
        "Room": (1660, 610, 1950, 730),
        "CourseOffering": (115, 940, 515, 1095),
        "Timetable": (115, 1220, 515, 1375),
        "ClassSession": (785, 1045, 1155, 1235),
        "TimeslotUse": (1375, 940, 1665, 1070),
        "RoomUse": (1375, 1185, 1665, 1315),
        "Policy": (785, 1325, 1155, 1460),
        "ChangeRequest": (1760, 1045, 2090, 1235),
        "Actors": (1760, 1325, 2090, 1460),
    }
    content = {
        "University": "name, campusName",
        "Faculty": "name",
        "Department": "name, code",
        "Teacher": "employeeNumber, maxDailyLoad",
        "StudentSection": "program, semester, batch, size",
        "Student": "registrationNumber",
        "Course": "code, title, creditHours",
        "CourseOffering": "links Course + Teacher + Section + Term",
        "AcademicTerm": "name, startDate, endDate",
        "Timeslot": "day, startTime, endTime",
        "Holiday": "name, date, reason",
        "Building": "name, location",
        "Room": "code, floor, capacity, type",
        "TimeslotUse": "same concept: Timeslot",
        "RoomUse": "same concept: Room",
        "Timetable": "status, score, conflicts",
        "ClassSession": "sessionNumber, status, locked",
        "Policy": "priority, weight, hard/soft",
        "ChangeRequest": "reason, urgency, status",
        "Actors": "Admin, Coordinator, Teacher, Student, Facility Manager",
    }
    for name, coords in boxes.items():
        fill = "#F8FBFD"
        if name in {"Timetable", "ClassSession", "ChangeRequest"}:
            fill = "#F7F9F2"
        elif name in {"Policy", "Holiday"}:
            fill = "#FFF6F6"
        display_name = {"TimeslotUse": "Timeslot", "RoomUse": "Room"}.get(name, name)
        box(draw, coords, display_name, content[name], fill=fill, outline=f"#{BLUE}")

    def mid_right(name):
        x1, y1, x2, y2 = boxes[name]
        return (x2, (y1 + y2) / 2)

    def mid_left(name):
        x1, y1, x2, y2 = boxes[name]
        return (x1, (y1 + y2) / 2)

    def mid_bottom(name):
        x1, y1, x2, y2 = boxes[name]
        return ((x1 + x2) / 2, y2)

    def mid_top(name):
        x1, y1, x2, y2 = boxes[name]
        return ((x1 + x2) / 2, y1)

    line(draw, mid_right("University"), mid_left("Faculty"), "contains 1..*")
    line(draw, mid_right("Faculty"), mid_left("Department"), "contains 1..*")
    line(draw, mid_right("Department"), mid_left("Teacher"), "employs 0..*")
    line(draw, mid_right("Department"), mid_left("StudentSection"), "contains 0..*")
    line(draw, mid_right("StudentSection"), mid_left("Student"), "contains 0..*")
    line(draw, mid_bottom("Department"), mid_top("Course"), "offers 0..*")
    line(draw, mid_right("AcademicTerm"), mid_left("Timeslot"), "defines 1..*")
    line(draw, (390, 735), mid_left("Holiday"), "records 0..*")
    line(draw, mid_right("Building"), mid_left("Room"), "houses 1..*")

    line(draw, mid_right("CourseOffering"), mid_left("ClassSession"), "realized-by 1..*")
    line(draw, mid_right("Timetable"), mid_left("ClassSession"), "contains 0..*")
    line(draw, mid_right("ClassSession"), mid_left("TimeslotUse"), "occurs-at 0..1")
    line(draw, mid_right("ClassSession"), mid_left("RoomUse"), "held-in 0..1")
    line(draw, mid_top("Policy"), mid_bottom("Timetable"), "guides 0..*")
    line(draw, mid_left("ChangeRequest"), mid_right("ClassSession"), "concerns 0..1")
    line(draw, mid_bottom("ChangeRequest"), mid_top("Actors"), "submitted/reviewed/decided by")

    draw.rounded_rectangle((70, 1510, 2130, 1570), radius=16, fill="#F5F5F5", outline="#DDDDDD")
    draw.text(
        (100, 1525),
        "Reading tip: some concepts appear once as master data and again as placement targets for readability.",
        font=font(25),
        fill=f"#{CHARCOAL}",
    )
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=CHARCOAL, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=9)
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row_data):
            cell = cells[c_idx]
            set_cell_text(cell, str(value), size=8.5)
            set_cell_margins(cell)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "FBFCFD")
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else GREEN)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(CHARCOAL)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    return p


def add_reference_image(doc, image_path, caption, width=6.4):
    if not image_path.exists():
        add_body(doc, f"Reference image not found: {image_path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    return p


def style_document(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Assignment 03 - UTOS Domain Model"
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("Assignment 03")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Domain Model for University Timetable Optimization and Management System")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(19)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project: SE-Hakari-Bankai / UTOS")
    r.font.name = "Calibri"
    r.font.size = Pt(12)

    details = [
        ("Instructor", "Sajid Anwar"),
        ("Submitted by", "Hakari Bankai Team"),
        ("Date", "May 9, 2026"),
        ("Assignment Requirement", "Develop a domain model showing key concepts, attributes, and relationships."),
    ]
    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    set_table_width(table, [1.9, 4.7])
    for row, (key, value) in zip(table.rows, details):
        set_cell_text(row.cells[0], key, bold=True, color=BLUE, size=10)
        set_cell_text(row.cells[1], value, size=10)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Domain model focuses on real-world concepts, not software design.")
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    method_img = ASSETS / "domain_model_method.png"
    overview_img = ASSETS / "utos_domain_model_overview.png"
    create_method_image(method_img)
    create_domain_diagram(overview_img)

    doc = Document()
    style_document(doc)
    add_cover(doc)
    doc.add_page_break()

    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "This document presents the domain model for the University Timetable Optimization and Management System (UTOS). "
        "It identifies the important real-world concepts in the project, their attributes, and the associations that connect them.",
    )
    add_body(
        doc,
        "The model follows object-oriented analysis principles: concepts are domain things such as Teacher, CourseOffering, Room, Timeslot, "
        "Timetable, and ChangeRequest. Software details such as database tables, APIs, screens, local storage, and algorithms are intentionally excluded.",
    )
    doc.add_picture(str(method_img), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "2. Concept Identification", 1)
    category_rows = [
        ("Physical/tangible objects", "Room, Building", "Rooms and buildings are real facilities used for classes."),
        ("Descriptions/specifications", "Course, CourseOffering, SchedulingPolicy", "Course describes subject; offering connects it to term, teacher, and section."),
        ("Places/containers", "University, Faculty, Department, StudentSection", "These contain people, courses, and academic groups."),
        ("Records/transactions", "Timetable, ChangeRequest", "These preserve scheduling decisions and requested changes."),
        ("Line items", "ClassSession", "One scheduled meeting inside a timetable."),
        ("Roles of people", "Teacher, Student, Administrator, Coordinator, FacilityManager", "Actors with real domain responsibility."),
        ("Rules/policies", "SchedulingPolicy, TeacherAvailability, Holiday", "Rules that guide or restrict timetable creation."),
    ]
    add_table(doc, ["Concept Category", "UTOS Concepts", "Reason"], category_rows, [1.65, 2.2, 2.75])

    add_heading(doc, "3. Key Concepts and Attributes", 1)
    add_body(
        doc,
        "Attributes are kept simple. Relationships are not stored as foreign-key attributes in the domain model; they are shown as associations.",
    )
    concept_rows = [
        ("University", "name, campusName"),
        ("Faculty", "name"),
        ("Department", "name, code"),
        ("Person", "name, email"),
        ("Teacher", "employeeNumber, maxDailyLoad"),
        ("Student", "registrationNumber"),
        ("TimetableAdministrator", "authorityLevel"),
        ("DepartmentCoordinator", "responsibilityArea"),
        ("FacilityManager", "responsibilityArea"),
        ("StudentSection", "name, program, semester, batch, size"),
        ("AcademicTerm", "name, startDate, endDate"),
        ("Course", "code, title, creditHours"),
        ("CourseOffering", "weeklySessions, requiredRoomType, sessionDuration"),
        ("Building", "name, location"),
        ("Room", "code, floor, capacity, roomType, features"),
        ("Timeslot", "day, startTime, endTime, isMorning, isLastSlot"),
        ("Holiday", "name, date, reason"),
        ("TeacherAvailability", "isAvailable, note"),
        ("SchedulingPolicy", "name, policyType, priority, weight, isHardConstraint, isEnabled"),
        ("Timetable", "name, status, qualityScore, hardConflictCount, softPenalty, unplacedSessionCount, createdAt"),
        ("ClassSession", "sessionNumber, status, locked"),
        ("ChangeRequest", "requestType, reason, urgency, preferredAlternative, coordinatorNote, administratorResponse, status, createdAt"),
    ]
    add_table(doc, ["Concept", "Attributes"], concept_rows, [2.05, 4.55])

    doc.add_page_break()
    add_heading(doc, "4. Domain Model Diagram", 1)
    add_body(
        doc,
        "The diagram below shows the main UTOS concepts and their need-to-know relationships. Multiplicity appears in the relationship labels.",
    )
    doc.add_picture(str(overview_img), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "5. Associations and Multiplicity", 1)
    association_rows = [
        ("University contains Faculty", "1 to 1..*", "Container relationship"),
        ("Faculty contains Department", "1 to 1..*", "Container relationship"),
        ("Department employs Teacher", "1 to 0..*", "Member/organization relationship"),
        ("Department contains StudentSection", "1 to 0..*", "Academic grouping"),
        ("Department offers Course", "1 to 0..*", "Course ownership"),
        ("StudentSection contains Student", "1 to 0..*", "Student membership"),
        ("AcademicTerm records CourseOffering", "1 to 0..*", "Offering belongs to a term"),
        ("Course is offered as CourseOffering", "1 to 0..*", "Description/specification separation"),
        ("Teacher teaches CourseOffering", "1 to 0..*", "Teaching assignment"),
        ("StudentSection takes CourseOffering", "1 to 0..*", "Section course load"),
        ("Building houses Room", "1 to 1..*", "Physical containment"),
        ("AcademicTerm defines Timeslot", "1 to 1..*", "Time structure"),
        ("AcademicTerm records Holiday", "1 to 0..*", "Blocked academic days"),
        ("Teacher states TeacherAvailability", "1 to 0..*", "Availability record"),
        ("TeacherAvailability refers to Timeslot", "1 to 1", "Availability is for a specific time"),
        ("Timetable contains ClassSession", "1 to 0..*", "ClassSession is a line item of Timetable"),
        ("CourseOffering is realized by ClassSession", "1 to 1..*", "Weekly sessions of one offering"),
        ("ClassSession occurs at Timeslot", "1 to 0..1", "0..1 allows unplaced sessions"),
        ("ClassSession is held in Room", "1 to 0..1", "0..1 allows unplaced sessions"),
        ("Timetable is guided by SchedulingPolicy", "1 to 0..*", "Rules and preferences"),
        ("Person submits ChangeRequest", "1 to 0..*", "Requested timetable change"),
        ("Coordinator reviews ChangeRequest", "1 to 0..*", "Review responsibility"),
        ("Administrator creates Timetable", "1 to 0..*", "Timetable management"),
        ("FacilityManager manages Room", "1 to 0..*", "Facility responsibility"),
        ("ChangeRequest concerns ClassSession/CourseOffering/Room/Timeslot", "1 to 0..1 target", "Request target"),
    ]
    add_table(doc, ["Association", "Multiplicity", "Purpose"], association_rows, [3.0, 1.15, 2.45])

    add_heading(doc, "6. Notes on Important Modeling Choices", 1)
    add_bullet(doc, "Course and CourseOffering are separate because a course description can be reused across terms, sections, and teachers.")
    add_bullet(doc, "ClassSession is separate because it is a line item inside a timetable and must remember placement, lock status, and scheduling status.")
    add_bullet(doc, "Room belongs to Building because room proximity and facility management depend on building/location information.")
    add_bullet(doc, "TeacherAvailability is modeled separately because availability is a relationship between a teacher and a timeslot with its own attribute.")
    add_bullet(doc, "SchedulingPolicy is included because hard constraints and soft preferences are part of the timetable domain, not just code.")
    add_bullet(doc, "Database IDs and foreign keys are not attributes in this domain model; associations carry those relationships.")

    doc.add_page_break()
    add_heading(doc, "7. Short Explanation", 1)
    add_body(
        doc,
        "The UTOS domain model represents a university scheduling environment. A university contains faculties and departments. "
        "Departments employ teachers, contain student sections, and offer courses. Students belong to sections, and sections take course offerings.",
    )
    add_body(
        doc,
        "A course offering connects a course to a teacher, section, and academic term. Each offering is realized by one or more class sessions. "
        "A timetable contains class sessions, and each class session may be assigned a timeslot and a room. If a session has no room or timeslot, it is still part of the timetable but remains unplaced.",
    )
    add_body(
        doc,
        "Scheduling policies, teacher availability, and holidays guide timetable validity and quality. Change requests allow teachers, students, coordinators, and administrators to communicate requested schedule changes. "
        "The model therefore captures the core concepts, attributes, and relationships needed for timetable generation, viewing, review, and improvement.",
    )

    add_heading(doc, "8. Excluded Software Artifacts", 1)
    excluded_rows = [
        ("Database / SQLite table", "Software storage design, not a real-world domain concept."),
        ("API endpoint / HTTP route", "Implementation mechanism."),
        ("Browser screen / form", "User interface artifact."),
        ("SolverRun / algorithm class", "Software execution detail."),
        ("teacherId, roomId, timeslotId", "Foreign-key style attributes; represented by associations instead."),
        ("CRUD operation", "Function/process, not a domain concept."),
    ]
    add_table(doc, ["Excluded Item", "Reason"], excluded_rows, [2.2, 4.4])

    doc.add_page_break()
    add_heading(doc, "9. Reviewed Project Diagrams", 1)
    add_body(
        doc,
        "These existing project diagrams were reviewed while building the domain model. They are included as reference material only. "
        "The final domain model removes database, software architecture, and process details where they are not real-world domain concepts.",
    )
    reference_rows = [
        ("Existing ERD/data model", "Useful for noun extraction: calendar, holidays, timetable, teachers, rooms, courses, sections, timeslots, users, change requests.", "Converted into domain concepts; PK/FK fields excluded."),
        ("Swimlane workflow", "Useful for use-case context: generation, review, publish, view, submit change request, approve/reject.", "Used to confirm ChangeRequest, Timetable, SchedulingPolicy, and actor roles."),
        ("Context diagram", "Useful for external actors: Admin, Teacher, Student, Department Coordinator, Facility Manager.", "Used to validate role concepts and major associations."),
        ("Architecture/DFD/process diagrams", "Useful for project background.", "Not part of core domain model because they describe software components or functions."),
    ]
    add_table(doc, ["Reviewed Diagram", "What It Helped With", "How It Was Used"], reference_rows, [1.65, 2.65, 2.3])

    add_reference_image(
        doc,
        SOURCE_DIAGRAMS / "Academic Timetable-2026-04-26-134435.png",
        "Reference Figure A: Existing ERD/data model used for concept extraction. Database keys and FK attributes are not carried into the final domain model.",
        width=6.5,
    )
    add_reference_image(
        doc,
        SOURCE_DIAGRAMS / "ChatGPT Image Apr 26, 2026, 08_46_14 PM.png",
        "Reference Figure B: Timetable generation workflow used to validate core actions and change request flow.",
        width=6.5,
    )
    add_reference_image(
        doc,
        SOURCE_DIAGRAMS / "Screenshot 2026-04-26 002329.png",
        "Reference Figure C: Context diagram used to validate major actor roles around UTOS.",
        width=6.3,
    )

    add_footer(doc)
    doc.core_properties.title = "Assignment 03 - UTOS Domain Model"
    doc.core_properties.subject = "Domain Model"
    doc.core_properties.author = "Hakari Bankai Team"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build_doc()
    print(out)
