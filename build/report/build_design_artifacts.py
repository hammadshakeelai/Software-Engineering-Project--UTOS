"""Build the supplementary design-artifacts document:
Packages + Package Diagram + CRC (Class-Responsibility-Collaborator) cards.
"""
import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from report_builder import (Report, TEAL, DARK, GREY, _set_cell_bg, _no_space,
                            COLW_IN)

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
GEN = os.path.join(HERE, "images", "generated")
AUTHOR = "Muhammad Hammad Shakeel"
COURSE = "Software Engineering"
EXTRA = ["Team Hakari Bankai", "4th Semester  ·  Spring 2026"]

PACKAGES = [
    ("Presentation", "«user interface»",
     "Vanilla HTML/CSS/JS single-page client: index.html, api.js, state.js, "
     "render.js, main.js. Renders role-based views and calls the Web API over JSON."),
    ("Web API", "«controller»",
     "server.py — a stdlib ThreadingHTTPServer that routes REST requests, "
     "enforces role permissions (X-User-Id), validates input, and returns JSON."),
    ("Application Services", "«service»",
     "bootstrap_service (initial payload) and timetable_service (solver "
     "orchestration + persistence). Coordinates repositories and the solver."),
    ("Solver", "«algorithm»",
     "timetable_solver — a constructive/backtracking scheduler behind a clean "
     "solve(problem) → result contract, swappable for OR-Tools CP-SAT."),
    ("Repositories", "«data access»",
     "master_data, timetable_repository, and system (notifications, audit). "
     "Encapsulate all SQL behind intention-revealing functions."),
    ("Domain Model", "«entities»",
     "The conceptual classes — Person hierarchy, CourseOffering, Timetable, "
     "ClassSession, Room, SchedulingPolicy, ChangeRequest, and so on."),
    ("Persistence", "«infrastructure»",
     "database.py, schema.sql, seed.py, migrate.py — SQLite connection, schema, "
     "seeding, and additive migrations run safely on startup."),
]

DEPS = [
    ["Client package", "depends on", "Server / responsibility"],
    ["Presentation", "→ calls", "Web API (JSON over HTTP)"],
    ["Web API", "→ delegates to", "Application Services, Solver"],
    ["Application Services", "→ uses", "Repositories, Solver, Domain Model"],
    ["Solver", "→ reads/writes", "Domain Model"],
    ["Repositories", "→ map", "Domain Model ↔ Persistence"],
    ["Persistence", "→ stores", "(no outgoing dependency — lowest layer)"],
]

# CRC cards: (Class, kind, [responsibilities], [collaborators])
CRC = [
    ("TimetableAdministrator", "control / actor",
     ["Generate, review, and publish official timetables",
      "Manage master data and lock fixed entries",
      "Approve, reject, and implement change requests"],
     ["Timetable", "TimetableService", "MasterData Repositories", "ChangeRequest"]),
    ("CourseOffering", "entity",
     ["Bind a Course to a Teacher, StudentSection, and AcademicTerm",
      "Know weekly session count and required room type",
      "Spawn the ClassSessions that must be scheduled"],
     ["Course", "Teacher", "StudentSection", "AcademicTerm", "ClassSession"]),
    ("Timetable", "entity",
     ["Hold all scheduled ClassSessions for a term",
      "Track status, score, hard conflicts, and soft penalty",
      "Support publish and export"],
     ["ClassSession", "TimetableVersion", "CourseOffering"]),
    ("ClassSession", "entity",
     ["Represent one placeable session of an offering",
      "Know its assigned Room and Timeslot and lock state",
      "Detect clashes with another session"],
     ["Room", "Timeslot", "CourseOffering", "Timetable"]),
    ("Timeslot", "entity",
     ["Define a weekly day + start/end time band",
      "Report duration and morning / last-slot flags"],
     ["ClassSession", "TeacherAvailability"]),
    ("Room", "entity",
     ["Know capacity, type, features, and building",
      "Report availability for a timeslot and utilization"],
     ["Building", "ClassSession"]),
    ("SchedulingPolicy", "entity",
     ["Hold hard constraints and weighted soft preferences",
      "Score a candidate placement and flag infeasibility"],
     ["Timetable", "ClassSession", "TimetableSolver"]),
    ("TeacherAvailability", "entity",
     ["Record whether a Teacher is free at a Timeslot"],
     ["Teacher", "Timeslot"]),
    ("ChangeRequest", "entity",
     ["Capture a requested change with reason and urgency",
      "Track status (pending/approved/rejected/implemented)",
      "Carry the coordinator recommendation note"],
     ["Teacher", "DepartmentCoordinator", "Timetable", "Notification"]),
    ("TimetableVersion", "entity",
     ["Snapshot a generated or repaired timetable",
      "Compute the diff (added/removed/changed) vs another version"],
     ["Timetable", "ClassSession"]),
    ("Notification", "entity",
     ["Deliver an event message to a recipient",
      "Track read/unread state"],
     ["Person", "ChangeRequest", "Timetable"]),
    ("AuditLog", "entity",
     ["Record every state-changing action",
      "Store actor, entity, and old/new values with a timestamp"],
     ["Person", "Repositories"]),
    ("TimetableService", "control",
     ["Build the solver problem and run the solver",
      "Persist the resulting version and reports",
      "Drive re-optimization preserving locked entries"],
     ["TimetableSolver", "TimetableRepository", "Timetable"]),
    ("TimetableSolver", "control / algorithm",
     ["Assign sessions to rooms and timeslots satisfying hard constraints",
      "Minimize weighted soft penalty",
      "Explain why each unplaced session could not be scheduled"],
     ["SchedulingPolicy", "ClassSession", "Room", "Timeslot"]),
]


def crc_card(rep, name, kind, resp, collab):
    doc = rep.doc
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = 1
    # row 0: class name (merged) with kind
    a = tbl.cell(0, 0).merge(tbl.cell(0, 1))
    _set_cell_bg(a, "0F6E66")
    p = a.paragraphs[0]; _no_space(p)
    r = p.add_run(name); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2 = p.add_run("    «%s»" % kind); r2.italic = True; r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0xD8, 0xEC, 0xE9)
    # row 1: column headers
    for ci, h in enumerate(["Responsibilities", "Collaborators"]):
        c = tbl.cell(1, ci); _set_cell_bg(c, "EAF3F1")
        p = c.paragraphs[0]; _no_space(p)
        rr = p.add_run(h); rr.bold = True; rr.font.size = Pt(9.5)
        rr.font.color.rgb = DARK
    # row 2: content
    rc = tbl.cell(2, 0); rc.text = ""
    for item in resp:
        p = rc.add_paragraph(); _no_space(p)
        run = p.add_run("•  " + item); run.font.size = Pt(9.5)
    cc = tbl.cell(2, 1); cc.text = ""
    for item in collab:
        p = cc.add_paragraph(); _no_space(p)
        run = p.add_run("•  " + item); run.font.size = Pt(9.5)
    tbl.columns[0].width = Inches(4.0)
    tbl.columns[1].width = Inches(2.5)
    for row in tbl.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(2.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def emit_packages(rep, loff=0):
    h = lambda lv: min(lv + loff, 4)
    rep.heading(h(1), "Logical Packages")
    rep.para("UTOS is organised into seven packages arranged as a strict "
             "dependency stack: higher layers depend only on the layers beneath "
             "them. This keeps coupling low (each package exposes a small "
             "interface) and cohesion high (each package has a single, clear "
             "responsibility).")
    for name, stereo, desc in PACKAGES:
        rep.heading(h(3), "%s  %s" % (name, stereo))
        rep.para(desc)
    rep.heading(h(2), "Package Dependencies")
    rep.table(DEPS, header=True, widths=[2.0, 1.3, 3.2])
    rep.doc.add_paragraph()
    rep.heading(h(2), "Package Diagram")
    rep.para("The figure below shows the seven packages as UML package (folder) "
             "symbols. Dashed arrows are dependency relationships («import / "
             "call»); every arrow points downward, confirming the acyclic "
             "layered structure.")
    rep.figure(os.path.join(GEN, "package_diagram.png"),
               "UTOS package diagram — layered architecture with «import/call» "
               "dependencies.", fig_no="P-1", max_h_in=6.5)


def emit_crc(rep, loff=0):
    rep.heading(min(1 + loff, 4), "CRC Cards")
    rep.para("Each card lists a class, its responsibilities (what it knows and "
             "does), and its collaborators (the classes it must work with). The "
             "cards were derived from the domain model and the class diagram and "
             "were used to validate that responsibilities are well distributed "
             "before detailed design.")
    rep.doc.add_paragraph()
    for name, kind, resp, collab in CRC:
        crc_card(rep, name, kind, resp, collab)


def build():
    rep = Report()
    rep.title_page(COURSE, "Design Artifacts",
                   "Packages, Package Diagram & CRC Cards",
                   "Logical Packaging and Class–Responsibility–Collaborator Analysis for UTOS",
                   AUTHOR, EXTRA)
    rep.page_break()
    rep.heading(1, "1. Introduction")
    rep.para("This document supplements the analysis and design artifacts with "
             "two further views of the UTOS architecture: a package diagram that "
             "groups the system's classes into cohesive, loosely-coupled layers, "
             "and CRC cards that summarise the responsibilities of the principal "
             "classes and the collaborators they rely on.")
    emit_packages(rep)
    rep.page_break()
    emit_crc(rep)
    rep.save(os.path.join(OUT, "Design_Artifacts_Packages_and_CRC.docx"))
    print("Design artifacts done")


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    build()
