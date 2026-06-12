"""Walk a .docx body in document order and emit a structured block list.

Blocks: heading / para / table / image, preserving order so a faithful
clean rebuild is possible. Images are written to images/<stem>/ and
referenced by filename. Run: python extract.py <docx> [<docx> ...]
"""
import sys, os, json, hashlib
from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_ROOT = os.path.join(HERE, "images")
OUT_ROOT = os.path.join(HERE, "extracted")


def style_name(p):
    try:
        return p.style.name or ""
    except Exception:
        return ""


def heading_level(p):
    s = style_name(p)
    if s.startswith("Heading"):
        try:
            return int(s.split()[-1])
        except Exception:
            return 1
    if s in ("Title",):
        return 0
    return None


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        yield child


def para_images(p, doc, img_dir, stem, counter):
    """Return list of saved image filenames found in this paragraph, in order."""
    out = []
    blips = p._p.findall('.//' + qn('a:blip'))
    for blip in blips:
        rid = blip.get(qn('r:embed')) or blip.get(qn('r:link'))
        if not rid:
            continue
        try:
            part = doc.part.related_parts[rid]
        except KeyError:
            continue
        blob = part.blob
        ext = os.path.splitext(part.partname)[1] or ".png"
        h = hashlib.md5(blob).hexdigest()[:8]
        counter[0] += 1
        fname = f"{stem}_img{counter[0]:02d}_{h}{ext}"
        fpath = os.path.join(img_dir, fname)
        if not os.path.exists(fpath):
            with open(fpath, "wb") as f:
                f.write(blob)
        out.append(fname)
    return out


def table_to_rows(tbl):
    rows = []
    for r in tbl.rows:
        rows.append([c.text for c in r.cells])
    return rows


def extract(path):
    doc = Document(path)
    stem = os.path.splitext(os.path.basename(path))[0]
    img_dir = os.path.join(IMG_ROOT, stem)
    os.makedirs(img_dir, exist_ok=True)
    os.makedirs(OUT_ROOT, exist_ok=True)

    # map xml element -> python object for paragraphs and tables
    para_map = {p._p: p for p in doc.paragraphs}
    tbl_map = {t._tbl: t for t in doc.tables}

    blocks = []
    counter = [0]
    for child in iter_block_items(doc):
        if child.tag == qn('w:p'):
            p = para_map.get(child)
            if p is None:
                continue
            imgs = para_images(p, doc, img_dir, stem, counter)
            txt = p.text
            lvl = heading_level(p)
            if imgs:
                for fn in imgs:
                    blocks.append({"type": "image", "file": fn})
            if txt.strip():
                if lvl is not None:
                    blocks.append({"type": "heading", "level": lvl, "text": txt})
                else:
                    blocks.append({"type": "para", "style": style_name(p), "text": txt})
        elif child.tag == qn('w:tbl'):
            t = tbl_map.get(child)
            if t is None:
                continue
            blocks.append({"type": "table", "rows": table_to_rows(t)})

    out = {"source": os.path.basename(path), "stem": stem,
           "images_dir": os.path.relpath(img_dir, HERE), "blocks": blocks}
    outpath = os.path.join(OUT_ROOT, stem + ".json")
    with open(outpath, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=1)
    nimg = sum(1 for b in blocks if b["type"] == "image")
    ntbl = sum(1 for b in blocks if b["type"] == "table")
    nhd = sum(1 for b in blocks if b["type"] == "heading")
    print(f"{stem}: blocks={len(blocks)} headings={nhd} tables={ntbl} images={nimg} -> {outpath}")


if __name__ == "__main__":
    for p in sys.argv[1:]:
        extract(p)
