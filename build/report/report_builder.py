"""Reusable clean-document builder for the UTOS SE assignments & final report.

Consistent academic styling: title page, numbered headings, banded tables,
captioned figures (auto-scaled to the text column). Built on python-docx.
"""
import os
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

TEAL = RGBColor(0x0F, 0x6E, 0x66)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x55, 0x5F, 0x6B)
LIGHT_FILL = "EAF3F1"
HEAD_FILL = "0F6E66"

# usable text column width (Letter, 1" margins)
COLW_IN = 6.5


def _set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


class Report:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        d = self.doc
        for s in d.sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)
        normal = d.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
        for i, sz in [(1, 16), (2, 13), (3, 11.5), (4, 10.5)]:
            st = d.styles["Heading %d" % i]
            st.font.name = "Calibri"
            st.font.size = Pt(sz)
            st.font.bold = True
            st.font.italic = (i == 4)
            st.font.color.rgb = TEAL if i == 1 else DARK
            st.paragraph_format.space_before = Pt(12 if i == 1 else 8)
            st.paragraph_format.space_after = Pt(4)
            st.paragraph_format.keep_with_next = True

    # ---- structural ----
    def page_break(self):
        self.doc.add_page_break()

    def title_page(self, course, assignment_no, title, subtitle, author,
                   extra_lines=None, instructor="Dr. Sajid Anwar"):
        d = self.doc
        for _ in range(3):
            d.add_paragraph()
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(course); r.font.size = Pt(13); r.font.color.rgb = GREY
        r.font.bold = True; r.font.name = "Calibri"
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(assignment_no); r.font.size = Pt(28); r.font.bold = True
        r.font.color.rgb = TEAL
        d.add_paragraph()
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title); r.font.size = Pt(20); r.font.bold = True
        r.font.color.rgb = DARK
        if subtitle:
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle); r.font.size = Pt(12.5); r.font.color.rgb = GREY
        for _ in range(6):
            d.add_paragraph()
        # rule
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("University Timetable Optimization & Management System (UTOS)")
        r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = TEAL
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Project: Optimize Flow  ·  Team Hakari Bankai")
        r.font.size = Pt(10.5); r.font.color.rgb = GREY
        d.add_paragraph()
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Submitted by"); r.font.size = Pt(10.5); r.font.color.rgb = GREY
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(author); r.font.size = Pt(13); r.font.bold = True
        r.font.color.rgb = DARK
        for line in (extra_lines or []):
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line); r.font.size = Pt(10.5); r.font.color.rgb = GREY
        if instructor:
            d.add_paragraph()
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("Submitted to"); r.font.size = Pt(10.5); r.font.color.rgb = GREY
            p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(instructor); r.font.size = Pt(12); r.font.bold = True
            r.font.color.rgb = DARK

    def heading(self, level, text):
        return self.doc.add_heading(text, level=level)

    def para(self, text, italic=False, bold=False, size=None, color=None,
             align=None, space_after=None):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.italic = italic; r.bold = bold
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        if align:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullet(self, text, level=0):
        p = self.doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.add_run(text)
        return p

    def numbered(self, text):
        return self.doc.add_paragraph(text, style="List Number")

    def callout(self, text, fill=LIGHT_FILL):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        _set_cell_bg(cell, fill)
        cell.width = Inches(COLW_IN)
        p = cell.paragraphs[0]
        r = p.add_run(text); r.font.size = Pt(10.5); r.italic = True
        r.font.color.rgb = DARK
        self.doc.add_paragraph()
        return tbl

    # ---- tables ----
    def table(self, rows, header=True, widths=None, font_size=9.5,
              first_col_bold=False):
        if not rows:
            return None
        ncol = max(len(r) for r in rows)
        tbl = self.doc.add_table(rows=len(rows), cols=ncol)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci in range(ncol):
                val = row[ci] if ci < len(row) else ""
                cell = tbl.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                _no_space(p)
                run = p.add_run(str(val))
                run.font.size = Pt(font_size)
                if header and ri == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_bg(cell, HEAD_FILL)
                else:
                    if ri % 2 == 0:
                        _set_cell_bg(cell, "F4F8F7")
                    if first_col_bold and ci == 0:
                        run.font.bold = True
        if widths:
            for ci, w in enumerate(widths):
                for ri in range(len(rows)):
                    tbl.cell(ri, ci).width = Inches(w)
        return tbl

    # ---- figures ----
    def figure(self, img_path, caption=None, width_in=None, max_h_in=8.2,
               fig_no=None):
        if not os.path.exists(img_path):
            self.para("[missing image: %s]" % os.path.basename(img_path),
                      italic=True, color=GREY)
            return
        try:
            iw, ih = Image.open(img_path).size
        except Exception:
            iw, ih = (1000, 700)
        w = width_in or COLW_IN
        h = w * ih / iw
        if h > max_h_in:
            h = max_h_in
            w = h * iw / ih
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(p)
        p.add_run().add_picture(img_path, width=Inches(w))
        if caption:
            cp = self.doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label = ("Figure %s. " % fig_no) if fig_no else ""
            r = cp.add_run(label + caption)
            r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GREY
            cp.paragraph_format.space_after = Pt(10)

    def toc_title(self, text="Table of Contents"):
        p = self.doc.add_paragraph()
        r = p.add_run(text); r.font.size = Pt(16); r.bold = True
        r.font.color.rgb = TEAL
        p.paragraph_format.space_after = Pt(8)

    def table_of_contents(self, depth="1-3"):
        """Insert a real Word TOC field (update with F9 / done via Word COM)."""
        p = self.doc.add_paragraph()
        run = p.add_run()
        fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "%s" \\h \\z \\u' % depth
        fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
        t = OxmlElement("w:t")
        t.text = "Update this field (select all, press F9) to build the contents."
        fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
        for el in (fb, instr, fs, t, fe):
            run._r.append(el)

    def save(self, path):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        self.doc.save(path)
        return path


# image-path resolver helpers --------------------------------------------------
def img(stem, filename):
    """Path to an extracted/raw image for a given source-doc stem."""
    p = os.path.join(HERE, "images", stem, filename)
    if os.path.exists(p):
        return p
    return os.path.join(HERE, "images", stem, "raw_media", filename)


def shot(name):
    return os.path.join(HERE, "screenshots", name + ".png")
